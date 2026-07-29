from fastapi import APIRouter, Request, File, UploadFile, Form
from fastapi.responses import HTMLResponse, StreamingResponse
import io
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from context import ctx

router = APIRouter()

@router.get("/", response_class=HTMLResponse)
async def pdf_home(request: Request):
    return request.app.state.templates.TemplateResponse("tool_pdf.html", ctx(request))

@router.post("/merge", response_class=StreamingResponse)
async def pdf_merge(request: Request, files: list[UploadFile] = File(...)):
    merger = PdfMerger()
    for f in files:
        merger.append(io.BytesIO(await f.read()))
    buf = io.BytesIO()
    merger.write(buf)
    merger.close()
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
                             headers={"Content-Disposition": "attachment; filename=merged.pdf"})


@router.post("/compress")
async def pdf_compress(request: Request, file: UploadFile = File(...)):
    reader = PdfReader(io.BytesIO(await file.read()))
    writer = PdfWriter()
    for page in reader.pages:
        page.compress_content_streams()
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
                             headers={"Content-Disposition": f"attachment; filename=compressed_{file.filename}"})


@router.post("/to-word")
async def pdf_to_word(request: Request, file: UploadFile = File(...)):
    """Convert PDF to Word (.docx) — extracts text from each page."""
    from docx import Document

    reader = PdfReader(io.BytesIO(await file.read()))
    doc = Document()
    doc.add_heading("Converted from PDF", level=1)

    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            doc.add_heading(f"Page {i}", level=2)
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    name = file.filename.rsplit(".", 1)[0] if file.filename else "converted"
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={name}.docx"})


@router.post("/split")
async def pdf_split(request: Request, file: UploadFile = File(...)):
    """Split PDF into individual pages as a ZIP."""
    import zipfile

    reader = PdfReader(io.BytesIO(await file.read()))
    buf = io.BytesIO()
    name = file.filename.rsplit(".", 1)[0] if file.filename else "split"

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(reader.pages, 1):
            writer = PdfWriter()
            writer.add_page(page)
            page_buf = io.BytesIO()
            writer.write(page_buf)
            page_buf.seek(0)
            zf.writestr(f"{name}_page_{i}.pdf", page_buf.read())

    buf.seek(0)
    return StreamingResponse(buf, media_type="application/zip",
                             headers={"Content-Disposition": f"attachment; filename={name}_split.zip"})


@router.post("/rotate")
async def pdf_rotate(request: Request, file: UploadFile = File(...), degrees: int = Form(90)):
    """Rotate all pages in a PDF by 90, 180, or 270 degrees."""
    reader = PdfReader(io.BytesIO(await file.read()))
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(degrees)
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
                             headers={"Content-Disposition": f"attachment; filename=rotated_{file.filename}"})


@router.get("/summarize", response_class=HTMLResponse)
async def pdf_summarize_page(request: Request):
    return request.app.state.templates.TemplateResponse(
        "tool_pdf.html", ctx(request, active_tab="summarize"))

@router.get("/chat", response_class=HTMLResponse)
async def pdf_chat_page(request: Request):
    return request.app.state.templates.TemplateResponse(
        "tool_pdf.html", ctx(request, active_tab="chat"))

@router.get("/to-word", response_class=HTMLResponse)
async def pdf_to_word_page(request: Request):
    return request.app.state.templates.TemplateResponse(
        "tool_pdf.html", ctx(request, active_tab="to-word"))

@router.get("/split", response_class=HTMLResponse)
async def pdf_split_page(request: Request):
    return request.app.state.templates.TemplateResponse(
        "tool_pdf.html", ctx(request, active_tab="split"))

@router.get("/rotate", response_class=HTMLResponse)
async def pdf_rotate_page(request: Request):
    return request.app.state.templates.TemplateResponse(
        "tool_pdf.html", ctx(request, active_tab="rotate"))
